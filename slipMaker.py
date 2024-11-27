import fitz  # PyMuPDF
import sys
import os
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import Tk
from tkinter.filedialog import askopenfilename, asksaveasfilename

def parse_pdf(file_path):
    """Extract lines from the PDF."""
    with fitz.open(file_path) as pdf:
        lines = []
        for page in pdf:
            lines.extend(page.get_text("text").splitlines())
    return lines

def process_loans(lines):
    """Process loan entries."""
    loans = []
    index = len(lines) - 1  # Start from the bottom and work backward
    while index >= 0:
        line = lines[index]
        if "Date retour prévue :" in line:
            loan = {}
            
            # Return Date
            loan["return_date"] = line.split(":")[1].strip()[:10]  # Extract YYYY-MM-DD
            index -= 1

            # Loan Date
            loan["loan_date"] = lines[index].split(":")[1].strip()[:10] if index >= 0 else ""
            index -= 1

            # Dewey
            loan["dewey"] = lines[index].strip() if index >= 0 else ""
            index -= 1

            # Identifier
            loan["identifier"] = lines[index].strip() if index >= 0 else ""
            index -= 1

            # Price (Ensure it starts with "Prix :")
            if index >= 0 and lines[index].startswith("Prix :"):
                loan["price"] = lines[index].strip()
                index -= 1
            else:
                loan["price"] = ""

            # Author
            if index >= 0 and "Prix :" not in lines[index]:
                loan["author"] = lines[index].strip()
                index -= 1
            else:
                loan["author"] = ""

            # Title (Accumulate until "Solde date du jour :")
            title_lines = []
            while index >= 0 and "Solde date du jour :" not in lines[index]:
                title_lines.insert(0, lines[index].strip())
                index -= 1
            loan["title"] = " ".join(title_lines).strip()

            # Handle case where title is blank and author is populated
            if not loan["title"] and loan["author"]:
                loan["title"] = loan["author"]
                loan["author"] = ""

            # Balance
            loan["balance"] = lines[index].split(":")[1].strip() if index >= 0 else ""
            index -= 1

            # Phone
            loan["phone"] = lines[index].replace("Téléphone :", "").strip() if index >= 0 else ""
            index -= 1

            # Address2
            loan["address2"] = lines[index].strip() if index >= 0 else ""
            index -= 1

            # Address1
            loan["address1"] = lines[index].strip() if index >= 0 else ""
            index -= 1

            # Name
            loan["name"] = lines[index].strip() if index >= 0 else ""
            index -= 1

            # Add the loan entry
            loans.append(loan)
        else:
            index -= 1  # Skip irrelevant lines

    return loans

def output_loans(loans):
    """Output the loans in the desired format."""
    # Group loans by borrower name
    borrower_loans = {}
    for loan in loans:
        name = loan["name"].strip()
        if name not in borrower_loans:
            borrower_loans[name] = []
        borrower_loans[name].append((loan["title"], loan["return_date"].replace("Date retour prévue :", "").strip()))
    
    # Generate output
    for borrower, books in borrower_loans.items():
        print(f"Borrower: {borrower}")
        for title, due_date in books:
            print(f"  Book: {title}")
            print(f"  Due Date: {due_date}")
        print("-" * 40)

def output_loans_by_loan(loans):
    """Prints all loan details in a formatted manner."""
    for loan in loans:
        print("--------------------------------------------------")
        print(f"Name: {loan['name']}")
        print(f"Address 1: {loan['address1']}")
        print(f"Address 2: {loan['address2']}")
        print(f"Phone: {loan['phone']}")
        print(f"Balance: {loan['balance']}")
        print(f"Title: {loan['title']}")
        print(f"Author: {loan['author']}")
        print(f"Dewey: {loan['dewey']}")
        print(f"Identifier: {loan['identifier']}")
        print(f"Price: {loan['price']}")
        print(f"Loan Date: {loan['loan_date']}")
        print(f"Return Date: {loan['return_date']}")
        print("--------------------------------------------------")

def flip_name_format(name):
    """Flip name format from LAST, FIRST to First Last."""
    if "," in name:
        last, first = map(str.strip, name.split(",", 1))
        return f"{first} {last}"
    return name  # Return as is if no comma

def resource_path(relative_path):
    """Get the absolute path to a resource, accounting for PyInstaller."""
    if getattr(sys, 'frozen', False):  # Running as a bundled app
        base_path = sys._MEIPASS  # Temporary folder created by PyInstaller
    else:
        base_path = os.path.abspath(".")  # Running in script mode
    return os.path.join(base_path, relative_path)

def create_landscape_two_column_slips(loans_by_borrower, logo_path=None, output_path="loans_slips.docx"):
    """Generate loan slips in landscape mode with two borrowers per page."""
    if logo_path is None:
        logo_path = resource_path("bookpic.png")
        
    doc = Document()

    # Set page orientation to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # Convert borrower dictionary to a list for iteration
    borrower_list = list(loans_by_borrower.items())
    num_borrowers = len(borrower_list)

    i = 0
    while i < num_borrowers:
        # Create a table for the page
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        col_width = (section.page_width - section.left_margin - section.right_margin) / 2

        # Add a row to the table
        row = table.add_row()
        left_cell, right_cell = row.cells

        # Populate the left cell with borrower info
        borrower, loans = borrower_list[i]
        populate_cell(left_cell, borrower, loans, logo_path, col_width)
        i += 1

        # Populate the right cell with borrower info if there's another borrower
        if i < num_borrowers:
            borrower, loans = borrower_list[i]
            populate_cell(right_cell, borrower, loans, logo_path, col_width)
            i += 1

        # Add a page break if there are more borrowers to process
        if i < num_borrowers:
            doc.add_page_break()

    # Save the document
    doc.save(output_path)
    print(f"Loan slips have been saved to {output_path}.")

def populate_cell(cell, borrower, loans, logo_path, col_width):
    """Populate a table cell with borrower information and loans."""
    cell.width = col_width
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add heading
    heading = cell_paragraph.add_run("📚 Livre parti à l'aventure!\n")
    heading.bold = True
    heading.font.size = Pt(16)
    heading.font.name = "Comic Sans MS"

    # Add Logo
    if logo_path:
        try:
            run = cell_paragraph.add_run()
            run.add_picture(logo_path, width=Cm(3))
            cell_paragraph.add_run("\n")
        except Exception as e:
            print(f"Error adding logo: {e}")

    # Add Borrower Name
    formatted_name = flip_name_format(borrower)
    borrower_paragraph = cell.add_paragraph()
    borrower_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    borrower_name = borrower_paragraph.add_run(f"Chez {formatted_name}\n")
    borrower_name.bold = True
    borrower_name.font.size = Pt(14)
    borrower_name.font.name = "Comic Sans MS"

    # Add playful message
    playful_message = cell.add_paragraph("⭐ Bonne Lecture! ⭐\n")
    playful_message.alignment = WD_ALIGN_PARAGRAPH.CENTER
    playful_message.style.font.name = "Comic Sans MS"
    playful_message.style.font.size = Pt(12)

    # Add Borrower's Loans
    for loan in loans:
        loan_title = loan['title']
        return_date = loan['return_date']

        # Truncate long titles
        max_title_length = 50
        if len(loan_title) > max_title_length:
            loan_title = loan_title[:max_title_length] + "..."

        # Compact format: Title and return date on one line
        loan_paragraph = cell.add_paragraph()
        loan_paragraph.add_run(f"📖 {loan_title} ").bold = True
        loan_paragraph.add_run(f"(Retour: {return_date})")
        loan_paragraph.style.font.name = "Comic Sans MS"
        loan_paragraph.style.font.size = Pt(10)

    # Add closing message
    closing_paragraph = cell.add_paragraph("\nMerci,\nLe comité bibliothèque")
    closing_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing_message = closing_paragraph.add_run()
    closing_message.bold = True
    closing_message.font.size = Pt(12)
    closing_message.font.name = "Comic Sans MS"


def get_input_file():
    """Prompt the user to select the input file."""
    root = Tk()
    root.withdraw()  # Hide the root window
    root.attributes('-topmost', True)  # Bring dialog to the front
    file_path = askopenfilename(
        title="Select the input PDF file",
        filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
    )
    root.destroy()
    return file_path

def get_output_file():
    """Prompt the user to select the output file location."""
    root = Tk()
    root.withdraw()  # Hide the root window
    root.attributes('-topmost', True)  # Bring dialog to the front
    file_path = asksaveasfilename(
        title="Save the output Word file",
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
    )
    root.destroy()
    return file_path

def main():
    """Main function to parse and process the PDF."""
    # Prompt for the input file
    file_path = get_input_file()
    if not file_path:
        print("No input file selected. Exiting.")
        return

    # Parse the input PDF
    lines = parse_pdf(file_path)
    loans = process_loans(lines)  # Process the extracted lines into structured loan data

    # Group loans by borrower name
    loans_by_borrower = {}
    for loan in loans:
        borrower = loan["name"]
        if borrower not in loans_by_borrower:
            loans_by_borrower[borrower] = []
        loans_by_borrower[borrower].append(loan)

    # Prompt for the output file
    output_path = get_output_file()
    if not output_path:
        print("No output file selected. Exiting.")
        return

    # Create the loan slips with cutlines
    create_landscape_two_column_slips(loans_by_borrower, output_path=output_path)
    

if __name__ == "__main__":
    main()
